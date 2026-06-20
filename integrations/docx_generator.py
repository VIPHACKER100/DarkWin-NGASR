"""DARKWIN DOCX Report Generator

Provides automated generation of professional Microsoft Word reports for security findings.
Supports custom branding, structured finding tables, and AI-assisted executive summaries.

Author: ARYAN AHIRWAR (VIPHACKER.100)
License: See LICENSE file
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from core.logging_system import get_logger

logger = get_logger("Integrations.DocxGenerator")


class DocxReportGenerator:
    """Generates professional DOCX reports from scan findings."""

    def __init__(self, target: str, scan_id: str) -> None:
        self.target = target
        self.scan_id = scan_id
        self.doc = Document()
        reports_dir = Path("reports")
        reports_dir.mkdir(exist_ok=True)
        self.report_path = str(reports_dir / f"DARKWIN_Report_{target}_{scan_id}.docx")

    def generate(self, findings: List[Dict[str, Any]], summary: str = "") -> str | None:
        """Create the DOCX report with findings and summary."""
        try:
            self._add_header()
            self._add_executive_summary(summary)
            self._add_findings_table(findings)
            self._add_footer()

            self.doc.save(self.report_path)
            logger.info(f"DOCX Report generated: {self.report_path}")
            return self.report_path
        except OSError as e:
            logger.error(f"Failed to generate DOCX report: {e}", exc_info=True)
            return None

    def _add_header(self):
        """Add report title and target information."""
        title = self.doc.add_heading('SECURITY ASSESSMENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Target: {self.target}\nScan ID: {self.scan_id}\nDate: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.bold = True

    def _add_executive_summary(self, summary: str):
        """Add AI-generated executive summary."""
        self.doc.add_heading('1. Executive Summary', level=1)
        if summary:
            self.doc.add_paragraph(summary)
        else:
            self.doc.add_paragraph("No executive summary provided for this scan.")

    def _add_findings_table(self, findings: List[Dict[str, Any]]):
        """Add a structured table of security findings."""
        self.doc.add_heading('2. Discovered Findings', level=1)
        
        if not findings:
            self.doc.add_paragraph("No vulnerabilities were identified during this assessment.")
            return

        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Endpoint'
        hdr_cells[3].text = 'Description'
        
        for f in findings:
            row_cells = table.add_row().cells
            row_cells[0].text = str(f.get('severity', 'N/A')).upper()
            row_cells[1].text = str(f.get('vuln_type', 'N/A'))
            row_cells[2].text = str(f.get('url', f.get('host', 'N/A')))
            row_cells[3].text = str(f.get('description', 'N/A'))

    def _add_footer(self):
        """Add developer attribution and legal disclaimer."""
        self.doc.add_page_break()
        self.doc.add_heading('3. Appendix & Legal', level=1)
        
        p = self.doc.add_paragraph()
        run = p.add_run("Developed by ARYAN AHIRWAR (VIPHACKER.100)")
        run.italic = True
        
        self.doc.add_paragraph(
            "Disclaimer: This report is for authorized security testing only. "
            "The developer is not responsible for any misuse or damage caused by this platform."
        )
